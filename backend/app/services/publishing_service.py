"""
출판/내보내기 서비스 모듈
DOCX, PDF, EPUB 형식으로 도서를 내보내는 기능을 제공합니다.
"""

import logging
import os
import subprocess
import tempfile
from typing import Any

from supabase import Client

from app.core.config import Settings
from app.models.base import TABLE_CHAPTERS, TABLE_EXPORTS
from app.schemas.design import PageSize
from app.schemas.publishing import ExportFormat, ExportStatusEnum

logger = logging.getLogger(__name__)


class PublishingService:
    """도서 출판/내보내기 서비스"""

    def __init__(self, settings: Settings) -> None:
        self._settings = settings
        self._export_dir = os.path.join(tempfile.gettempdir(), "moduga_exports")
        os.makedirs(self._export_dir, exist_ok=True)

    async def start_export(
        self,
        export_id: str,
        book_data: dict[str, Any],
        export_format: ExportFormat,
        page_size: PageSize = PageSize.A5,
        include_cover: bool = True,
        include_toc: bool = True,
        accessibility_tags: bool = True,
        supabase: Client | None = None,
    ) -> str:
        """
        도서 내보내기를 시작합니다.

        Args:
            export_id: 내보내기 ID
            book_data: 도서 정보
            export_format: 내보내기 형식
            page_size: 페이지 크기
            include_cover: 표지 포함 여부
            include_toc: 목차 포함 여부
            accessibility_tags: 접근성 태그 포함 여부 (EPUB)
            supabase: Supabase 클라이언트

        Returns:
            생성된 파일 경로
        """
        if not supabase:
            raise ValueError("Supabase 클라이언트가 필요합니다.")

        # 챕터 데이터 조회
        chapters_resp = (
            supabase.table(TABLE_CHAPTERS)
            .select("*")
            .eq("book_id", book_data["id"])
            .order("order", desc=False)
            .execute()
        )
        chapters = chapters_resp.data

        # 진행 상태 업데이트: 처리 중
        supabase.table(TABLE_EXPORTS).update({
            "status": ExportStatusEnum.PROCESSING.value,
            "progress": 10.0,
        }).eq("id", export_id).execute()

        try:
            # 형식별 내보내기 처리
            if export_format == ExportFormat.DOCX:
                file_path = await self._export_docx(
                    export_id, book_data, chapters, include_toc
                )
            elif export_format == ExportFormat.PDF:
                file_path = await self._export_pdf(
                    export_id, book_data, chapters, page_size, include_toc
                )
            elif export_format == ExportFormat.EPUB:
                file_path = await self._export_epub(
                    export_id, book_data, chapters, include_toc, accessibility_tags
                )
            else:
                raise ValueError(f"지원하지 않는 형식: {export_format}")

            # 파일 크기 확인
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0

            # 완료 상태 업데이트
            supabase.table(TABLE_EXPORTS).update({
                "status": ExportStatusEnum.COMPLETED.value,
                "progress": 100.0,
                "file_path": file_path,
                "file_size_bytes": file_size,
            }).eq("id", export_id).execute()

            return file_path

        except Exception as e:
            # 실패 상태 업데이트
            supabase.table(TABLE_EXPORTS).update({
                "status": ExportStatusEnum.FAILED.value,
                "error_message": str(e),
            }).eq("id", export_id).execute()
            raise

    async def _export_docx(
        self,
        export_id: str,
        book_data: dict[str, Any],
        chapters: list[dict[str, Any]],
        include_toc: bool = True,
    ) -> str:
        """
        DOCX 형식으로 내보냅니다.

        Args:
            export_id: 내보내기 ID
            book_data: 도서 정보
            chapters: 챕터 데이터 목록
            include_toc: 목차 포함 여부

        Returns:
            생성된 DOCX 파일 경로
        """
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 문서 스타일 설정
        style = doc.styles["Normal"]
        font = style.font
        font.name = "맑은 고딕"
        font.size = Pt(11)

        # 제목 페이지
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(book_data.get("title", "제목 없음"))
        title_run.font.size = Pt(24)
        title_run.bold = True

        doc.add_paragraph()  # 빈 줄

        # 부제 / 설명
        description = book_data.get("description", "")
        if description:
            desc_para = doc.add_paragraph()
            desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc_para.add_run(description).font.size = Pt(12)

        doc.add_page_break()

        # 목차 (간단한 텍스트 형태)
        if include_toc and chapters:
            toc_heading = doc.add_heading("목차", level=1)
            for ch in chapters:
                toc_entry = doc.add_paragraph(
                    f"{ch.get('order', 0)}. {ch.get('title', '')}",
                    style="List Number",
                )
            doc.add_page_break()

        # 챕터 내용
        for ch in chapters:
            # 챕터 제목
            doc.add_heading(ch.get("title", ""), level=1)

            # 챕터 본문
            content = ch.get("content", "")
            if content:
                paragraphs = content.split("\n")
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        doc.add_paragraph(para_text)

            doc.add_page_break()

        # 파일 저장
        file_path = os.path.join(self._export_dir, f"{export_id}.docx")
        doc.save(file_path)

        logger.info(f"DOCX 생성 완료: {file_path}")
        return file_path

    async def _export_pdf(
        self,
        export_id: str,
        book_data: dict[str, Any],
        chapters: list[dict[str, Any]],
        page_size: PageSize = PageSize.A5,
        include_toc: bool = True,
    ) -> str:
        """
        PDF 형식으로 내보냅니다. Typst를 사용하여 조판합니다.

        Args:
            export_id: 내보내기 ID
            book_data: 도서 정보
            chapters: 챕터 데이터 목록
            page_size: 페이지 크기
            include_toc: 목차 포함 여부

        Returns:
            생성된 PDF 파일 경로
        """
        # 페이지 크기 매핑
        page_size_map: dict[str, str] = {
            PageSize.A5.value: "a5",
            PageSize.B5.value: "b5",
            PageSize.A4.value: "a4",
            PageSize.PAPERBACK.value: "us-letter",
        }
        typst_page_size = page_size_map.get(page_size.value, "a5")

        title = book_data.get("title", "제목 없음")

        # Typst 소스 생성
        typst_parts = [
            f'#set page(paper: "{typst_page_size}", margin: (top: 25mm, bottom: 25mm, left: 20mm, right: 15mm))',
            '#set text(font: "Noto Sans CJK KR", size: 11pt, lang: "ko")',
            "#set par(leading: 0.8em, justify: true)",
            "",
            "// 제목 페이지",
            "#align(center + horizon)[",
            f"  #text(size: 24pt, weight: \"bold\")[{self._escape_typst(title)}]",
            "]",
            "#pagebreak()",
        ]

        # 목차
        if include_toc:
            typst_parts.extend([
                "// 목차",
                "#outline(indent: 1em)",
                "#pagebreak()",
            ])

        # 챕터 내용
        for ch in chapters:
            ch_title = self._escape_typst(ch.get("title", ""))
            ch_content = self._escape_typst(ch.get("content", ""))

            typst_parts.extend([
                f"= {ch_title}",
                "",
                ch_content,
                "",
                "#pagebreak()",
            ])

        typst_source = "\n".join(typst_parts)

        # Typst 컴파일
        src_path = os.path.join(self._export_dir, f"{export_id}.typ")
        pdf_path = os.path.join(self._export_dir, f"{export_id}.pdf")

        with open(src_path, "w", encoding="utf-8") as f:
            f.write(typst_source)

        try:
            result = subprocess.run(
                ["typst", "compile", src_path, pdf_path],
                capture_output=True,
                text=True,
                timeout=60,
            )

            if result.returncode != 0:
                raise RuntimeError(f"Typst 컴파일 오류: {result.stderr}")

            logger.info(f"PDF 생성 완료: {pdf_path}")
            return pdf_path

        except FileNotFoundError:
            raise RuntimeError(
                "Typst가 설치되어 있지 않습니다. "
                "https://typst.app 에서 설치해주세요."
            )

    async def _export_epub(
        self,
        export_id: str,
        book_data: dict[str, Any],
        chapters: list[dict[str, Any]],
        include_toc: bool = True,
        accessibility_tags: bool = True,
    ) -> str:
        """
        EPUB 형식으로 내보냅니다. 접근성 태그를 포함합니다.

        Args:
            export_id: 내보내기 ID
            book_data: 도서 정보
            chapters: 챕터 데이터 목록
            include_toc: 목차 포함 여부
            accessibility_tags: 접근성 태그 포함 여부

        Returns:
            생성된 EPUB 파일 경로
        """
        from ebooklib import epub

        book = epub.EpubBook()

        # 메타데이터 설정
        title = book_data.get("title", "제목 없음")
        book.set_identifier(f"moduga-jakga-{export_id}")
        book.set_title(title)
        book.set_language("ko")
        book.add_author(book_data.get("author_name", "작가"))

        # 접근성 메타데이터 (EPUB Accessibility 규격)
        if accessibility_tags:
            book.add_metadata(
                "http://purl.org/dc/terms/",
                "conformsTo",
                "EPUB Accessibility 1.1 - WCAG 2.1 Level AA",
            )
            book.add_metadata(
                "http://www.idpf.org/2007/opf",
                "meta",
                "시각장애인 접근성 지원",
                {"property": "schema:accessibilitySummary"},
            )
            book.add_metadata(
                "http://www.idpf.org/2007/opf",
                "meta",
                "textual",
                {"property": "schema:accessMode"},
            )
            book.add_metadata(
                "http://www.idpf.org/2007/opf",
                "meta",
                "textual",
                {"property": "schema:accessModeSufficient"},
            )

        # CSS 스타일
        style = """
        body {
            font-family: "맑은 고딕", "Malgun Gothic", sans-serif;
            font-size: 1em;
            line-height: 1.8;
            text-align: justify;
        }
        h1 {
            font-size: 1.5em;
            text-align: center;
            margin-top: 2em;
            margin-bottom: 1em;
        }
        p {
            margin-bottom: 0.5em;
            text-indent: 1em;
        }
        """
        nav_css = epub.EpubItem(
            uid="style_nav",
            file_name="style/nav.css",
            media_type="text/css",
            content=style.encode("utf-8"),
        )
        book.add_item(nav_css)

        # 챕터 생성
        epub_chapters: list[epub.EpubHtml] = []
        spine_items: list[Any] = ["nav"]

        for ch in chapters:
            ch_title = ch.get("title", "")
            ch_content = ch.get("content", "")
            ch_order = ch.get("order", 1)

            # HTML 변환
            html_paragraphs = ""
            if ch_content:
                for para in ch_content.split("\n"):
                    para = para.strip()
                    if para:
                        html_paragraphs += f"<p>{self._escape_html(para)}</p>\n"

            # 접근성 속성 추가
            role_attr = ' role="doc-chapter"' if accessibility_tags else ""
            aria_label = f' aria-label="챕터 {ch_order}: {ch_title}"' if accessibility_tags else ""

            html_content = f"""
            <html xmlns="http://www.w3.org/1999/xhtml">
            <head><link rel="stylesheet" href="style/nav.css" type="text/css"/></head>
            <body>
            <section{role_attr}{aria_label}>
                <h1>{self._escape_html(ch_title)}</h1>
                {html_paragraphs}
            </section>
            </body>
            </html>
            """

            epub_ch = epub.EpubHtml(
                title=ch_title,
                file_name=f"chapter_{ch_order:03d}.xhtml",
                lang="ko",
            )
            epub_ch.content = html_content.encode("utf-8")
            epub_ch.add_item(nav_css)

            book.add_item(epub_ch)
            epub_chapters.append(epub_ch)
            spine_items.append(epub_ch)

        # 목차 설정
        book.toc = epub_chapters

        # Navigation 파일 추가
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # 스파인 (읽는 순서) 설정
        book.spine = spine_items

        # EPUB 파일 저장
        file_path = os.path.join(self._export_dir, f"{export_id}.epub")
        epub.write_epub(file_path, book)

        logger.info(f"EPUB 생성 완료: {file_path}")
        return file_path

    def _escape_typst(self, text: str) -> str:
        """Typst 특수 문자를 이스케이프합니다."""
        # Typst에서 의미가 있는 특수 문자 이스케이프
        replacements = {
            "#": "\\#",
            "$": "\\$",
            "*": "\\*",
            "_": "\\_",
            "@": "\\@",
            "<": "\\<",
            ">": "\\>",
        }
        for char, escaped in replacements.items():
            text = text.replace(char, escaped)
        return text

    def _escape_html(self, text: str) -> str:
        """HTML 특수 문자를 이스케이프합니다."""
        replacements = {
            "&": "&amp;",
            "<": "&lt;",
            ">": "&gt;",
            '"': "&quot;",
            "'": "&#x27;",
        }
        for char, escaped in replacements.items():
            text = text.replace(char, escaped)
        return text
